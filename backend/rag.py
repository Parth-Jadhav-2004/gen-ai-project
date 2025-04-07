import numpy as np
from vectorstore import build_index, get_embedding
from mistral_client import get_answer
from PyPDF2 import PdfReader
from docx import Document
import ollama
import io

def extract_text_from_pdf(file_bytes):
    reader = PdfReader(io.BytesIO(file_bytes))
    return " ".join(
        page.extract_text() for page in reader.pages if page.extract_text()
    )

def extract_text_from_docx(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join(
        [para.text for para in doc.paragraphs if para.text.strip()]
    )

def chunk_text(text, page_number=None):
    chunk_size = 1000
    chunk_overlap = 200
    chunks = []

    for i in range(0, len(text), chunk_size - chunk_overlap):
        chunk = {
            "text": text[i:i + chunk_size],
            "page_number": page_number
        }
        chunks.append(chunk)

    return chunks

def process_documents(documents, filenames):
    all_chunks = []

    for content, filename in zip(documents, filenames):
        if filename.endswith(".pdf"):
            reader = PdfReader(io.BytesIO(content))
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    chunks = chunk_text(page_text, page_number=i + 1)
                    all_chunks.extend(chunks)

        elif filename.endswith(".docx"):
            doc = Document(io.BytesIO(content))
            full_text = "\n".join(
                [para.text for para in doc.paragraphs if para.text.strip()]
            )
            chunks = chunk_text(full_text)
            all_chunks.extend(chunks)

        else:  # assume .txt
            text = content.decode('utf-8')
            chunks = chunk_text(text)
            all_chunks.extend(chunks)

    # ✅ Ensure only string texts go to build_index
    chunk_texts = [chunk["text"] for chunk in all_chunks if isinstance(chunk["text"], str)]

    # 🔐 Sanity check
    assert all(isinstance(text, str) for text in chunk_texts), "Non-string found in chunk_texts"

    index = build_index(chunk_texts)

    return all_chunks, index

def answer_question(index, chunks, question):
    q_embedding = get_embedding(question)
    scores = np.dot(index["embeddings"], q_embedding)
    top_k = scores.argsort()[-1:][::-1]  # Use only top-1 to pinpoint page

    best_chunk = chunks[top_k[0]]
    context = best_chunk["text"]
    page_number = best_chunk["page_number"]

    answer = get_answer(context, question)

    return {
        "answer": answer,
        "page_number": page_number
    }

def summarize_chunks(chunks):
    context = "\n".join(chunk["text"] for chunk in chunks)
    prompt = f"Summarize the following document:\n\n{context}"

    response = ollama.chat(
        model="mistral",
        messages=[
            {"role": "system", "content": "You are a summarizer."},
            {"role": "user", "content": prompt}
        ]
    )
    summary_text = response["message"]["content"]
    
    # Ensure we return a list for frontend
    return [point.strip() for point in summary_text.split('\n') if point.strip()]
